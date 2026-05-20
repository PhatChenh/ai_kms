"""DOCX handler — extracts paragraph text from .docx drops.

Uses python-docx to iterate paragraphs. Empty DOCX files (all paragraphs
blank) return Success with text="" — empty text is valid input for the LLM
summarise stage, which handles the no-content case.

TD-H1: Table cell text is not extracted. python-docx exposes doc.tables but
iterating them is deferred to Phase 3+. Document the limitation here so the
LLM stage can flag notes with tables as potentially incomplete.
"""
from pathlib import Path

from docx import Document

from handlers.base import BaseHandler, RawContent
from handlers.registry import HandlerRegistry
from core.result import Failure, Result, Success

__all__ = ["DocxHandler"]


@HandlerRegistry.register
class DocxHandler(BaseHandler):
    """Handles .docx (Word) files dropped into the vault inbox.

    Extracts non-empty paragraph text only. Table cells are not included
    (TD-H1 — deferred to Phase 3+).
    """

    def can_handle(self, path: Path) -> bool:
        """Return True for .docx files (case-insensitive).

        Args:
            path: Path to the dropped file.

        Returns:
            True if the file extension is .docx regardless of case.
        """
        return path.suffix.lower() == ".docx"

    def extract(self, path: Path) -> Result[RawContent]:
        """Extract paragraph text from a DOCX file.

        Args:
            path: Path to the .docx file.

        Returns:
            Success(RawContent) with paragraph text joined by newlines, is_md=False.
            Empty DOCX returns Success with text="" (not Failure).
            Failure(recoverable=False) if python-docx raises any exception
            (corrupt file, missing file, not a DOCX, etc.).
        """
        try:
            doc = Document(str(path))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as exc:
            return Failure(
                error=f"DOCX read failed: {exc}",
                recoverable=False,
                context={"path": str(path)},
            )

        return Success(RawContent(text=text, source_path=path, is_md=False))
