"""DOCX handler — extracts paragraph text from .docx drops.

Uses python-docx to iterate paragraphs. Empty DOCX files (all paragraphs
blank) return Success with text="" — empty text is valid input for the LLM
summarise stage, which handles the no-content case.

TD-H1: Table cell text is not extracted. python-docx exposes doc.tables but
iterating them is deferred to Phase 3+. Document the limitation here so the
LLM stage can flag notes with tables as potentially incomplete.

Files larger than CONFIG.main.handlers.max_file_size_bytes are rejected before
any parse work to keep a single drop from OOMing the capture process.
"""
from pathlib import Path

import structlog
from docx import Document

from handlers.base import BaseHandler, RawContent
from handlers.registry import HandlerRegistry
from core.result import Failure, Result, Success

__all__ = ["DocxHandler"]

logger = structlog.get_logger(__name__)


@HandlerRegistry.register
class DocxHandler(BaseHandler):
    """Handles .docx (Word) files dropped into the vault inbox.

    Extracts non-empty paragraph text only. Table cells are not included
    (TD-H1 — deferred to Phase 3+).
    """

    def can_handle(self, path: Path) -> bool:
        return path.suffix.lower() == ".docx"

    def extract(self, path: Path, *, max_file_size_bytes: int | None = None) -> Result[RawContent]:
        if max_file_size_bytes is None:
            # Lazy import — see handlers/pdf_handler.py for rationale.
            from core.config import CONFIG

            max_bytes = CONFIG.main.handlers.max_file_size_bytes
        else:
            max_bytes = max_file_size_bytes

        try:
            size = path.stat().st_size
        except OSError as exc:
            logger.warning("docx.stat.failed", path=str(path), error=str(exc))
            return Failure(
                error=f"DOCX stat failed: {exc}",
                recoverable=False,
                context={"path": str(path)},
            )

        if size > max_bytes:
            logger.warning(
                "docx.too_large", path=str(path), size=size, limit=max_bytes
            )
            return Failure(
                error=f"DOCX too large: {size} > {max_bytes} bytes",
                recoverable=False,
                context={"path": str(path), "size": size, "limit": max_bytes},
            )

        try:
            doc = Document(str(path))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as exc:
            logger.warning("docx.read.failed", path=str(path), error=str(exc))
            return Failure(
                error=f"DOCX read failed: {exc}",
                recoverable=False,
                context={"path": str(path)},
            )

        logger.info(
            "docx.extract.ok", path=str(path), chars=len(text), bytes=size
        )
        return Success(RawContent(text=text, source_path=path, is_md=False))
