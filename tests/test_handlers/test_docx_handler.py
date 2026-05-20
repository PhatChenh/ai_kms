"""Tests for handlers/docx_handler.py — DocxHandler end-to-end."""
from pathlib import Path

import pytest
from docx import Document

from handlers.docx_handler import DocxHandler
from handlers.base import RawContent
from core.result import Failure, Success


@pytest.fixture
def docx_file(tmp_path: Path) -> Path:
    path = tmp_path / "note.docx"
    doc = Document()
    doc.add_paragraph("Hello World")
    doc.add_paragraph("Second paragraph")
    doc.save(str(path))
    return path


@pytest.fixture
def empty_docx_file(tmp_path: Path) -> Path:
    path = tmp_path / "empty.docx"
    doc = Document()
    doc.save(str(path))
    return path


def test_extract_returns_success_with_raw_content(docx_file: Path):
    result = DocxHandler().extract(docx_file)
    assert isinstance(result, Success)
    assert isinstance(result.value, RawContent)


def test_extract_text_contains_hello_world(docx_file: Path):
    result = DocxHandler().extract(docx_file)
    assert isinstance(result, Success)
    assert "Hello World" in result.value.text


def test_extract_text_contains_all_paragraphs(docx_file: Path):
    result = DocxHandler().extract(docx_file)
    assert isinstance(result, Success)
    assert "Second paragraph" in result.value.text


def test_extract_sets_is_md_false(docx_file: Path):
    result = DocxHandler().extract(docx_file)
    assert isinstance(result, Success)
    assert result.value.is_md is False


def test_extract_sets_source_path(docx_file: Path):
    result = DocxHandler().extract(docx_file)
    assert isinstance(result, Success)
    assert result.value.source_path == docx_file


def test_empty_docx_returns_success_with_empty_text(empty_docx_file: Path):
    result = DocxHandler().extract(empty_docx_file)
    assert isinstance(result, Success)
    assert result.value.text == ""


def test_nonexistent_path_returns_failure(tmp_path: Path):
    result = DocxHandler().extract(tmp_path / "missing.docx")
    assert isinstance(result, Failure)
    assert result.recoverable is False


def test_can_handle_docx_lowercase():
    assert DocxHandler().can_handle(Path("doc.docx")) is True


def test_can_handle_docx_uppercase():
    assert DocxHandler().can_handle(Path("DOC.DOCX")) is True


def test_can_handle_rejects_pdf():
    assert DocxHandler().can_handle(Path("doc.pdf")) is False
